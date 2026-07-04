#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
翻訳ウォッチャー Pro
テキスト / Excel / Word ファイルを自動翻訳
Google翻訳 / ローカルLLM（Ollama）対応  |  Windows・macOS 両対応
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading, time, os, sys, json, platform, urllib.request, urllib.error
from pathlib import Path
from datetime import datetime

# ── PyInstaller対応 ──────────────────────────────────
def _base() -> Path:
    return Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent

BASE_DIR    = _base()
CONFIG_FILE = BASE_DIR / "config.json"
IS_WIN      = platform.system() == "Windows"
FONT_UI     = "Yu Gothic UI" if IS_WIN else "Hiragino Kaku Gothic ProN"
FONT_MONO   = "MS Gothic"    if IS_WIN else "Courier New"
APP_RED     = "#b91c1c"
APP_DARK    = "#1e1b1b"
APP_BG      = "#f8f8f8"
APP_PANEL   = "#ffffff"

def F(sz=11, bold=False): return (FONT_UI, sz, "bold" if bold else "normal")

# ── オプションライブラリ ──────────────────────────────
try:
    from deep_translator import GoogleTranslator; GOOGLE_OK = True
except ImportError: GOOGLE_OK = False

try:
    import openpyxl; EXCEL_OK = True
except ImportError: EXCEL_OK = False

try:
    from docx import Document as _Docx; WORD_OK = True
except ImportError: WORD_OK = False

try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler; WATCH_OK = True
except ImportError: WATCH_OK = False

EXTS = {".txt"} | ({".xlsx",".xls"} if EXCEL_OK else set()) | ({".docx",".doc"} if WORD_OK else set())
LANGS = {"vi": "ベトナム語", "zh-CN": "中国語（簡体）"}

# ══════════════════════════════════════════════════════
class App:
    def __init__(self, root: tk.Tk):
        self.root    = root
        self.root.title("翻訳ウォッチャー Pro")
        self.root.geometry("960x660")
        self.root.minsize(720, 520)
        self.root.configure(bg=APP_BG)
        self.cfg      = self._cfg_load()
        self.observer = None
        self.watching = False
        self.stats    = {"total":0, "done":0, "error":0}
        self.in_dir   = Path(self.cfg.get("in_dir",  str(BASE_DIR/"input")))
        self.out_dir  = Path(self.cfg.get("out_dir", str(BASE_DIR/"output")))
        self.in_dir.mkdir(exist_ok=True)
        self.out_dir.mkdir(exist_ok=True)
        self._style()
        self._menu()
        self._toolbar()
        self._body()
        self._statusbar()
        self._boot_log()

    # ── スタイル ────────────────────────────────────
    def _style(self):
        s = ttk.Style()
        s.theme_use("clam")
        s.configure(".", font=F(11), background=APP_BG)
        s.configure("Toolbar.TFrame", background="#2d2d2d")
        s.configure("Toolbar.TLabel", background="#2d2d2d", foreground="#ffffff", font=F(11))
        s.configure("Toolbar.TButton",
                    background="#444444", foreground="#ffffff",
                    relief="flat", padding=(10,5), font=F(11))
        s.map("Toolbar.TButton",
              background=[("active","#666666"),("pressed","#b91c1c")])
        s.configure("Red.TButton",
                    background=APP_RED, foreground="#ffffff",
                    relief="flat", padding=(12,5), font=F(11,bold=True))
        s.map("Red.TButton", background=[("active","#991b1b")])
        s.configure("Card.TFrame",  background=APP_PANEL, relief="solid", borderwidth=1)
        s.configure("Card.TLabel",  background=APP_PANEL)
        s.configure("Stat.TLabel",  background=APP_PANEL, font=F(11))
        s.configure("Big.TLabel",   background=APP_PANEL, font=F(26,bold=True))
        s.configure("Status.TLabel",background="#2d2d2d", foreground="#cccccc", font=F(10))
        s.configure("Treeview", font=F(11), rowheight=26,
                    background=APP_PANEL, fieldbackground=APP_PANEL)
        s.configure("Treeview.Heading", font=F(11,bold=True),
                    background="#e8e8e8", relief="flat")
        s.map("Treeview", background=[("selected","#dbeafe")])

    # ── メニューバー ─────────────────────────────────
    def _menu(self):
        mb = tk.Menu(self.root)
        self.root.config(menu=mb)

        fm = tk.Menu(mb, tearoff=0)
        fm.add_command(label="ファイルを追加...",    command=self._add_files)
        fm.add_command(label="入力フォルダを開く",   command=lambda: _open(self.in_dir))
        fm.add_command(label="出力フォルダを開く",   command=lambda: _open(self.out_dir))
        fm.add_separator()
        fm.add_command(label="リストをクリア",       command=self._clear_list)
        fm.add_separator()
        fm.add_command(label="終了",                command=self.on_close)
        mb.add_cascade(label="ファイル(F)", menu=fm)

        sm = tk.Menu(mb, tearoff=0)
        sm.add_command(label="設定...", command=self._open_settings)
        mb.add_cascade(label="設定(S)", menu=sm)

        hm = tk.Menu(mb, tearoff=0)
        hm.add_command(label="使い方", command=self._show_help)
        mb.add_cascade(label="ヘルプ(H)", menu=hm)

    # ── ツールバー ───────────────────────────────────
    def _toolbar(self):
        tb = tk.Frame(self.root, bg="#2d2d2d", pady=6)
        tb.pack(fill="x")

        self.watch_btn = tk.Button(tb, text="▶  監視 開始",
            bg="#166534", fg="white", relief="flat",
            padx=14, pady=5, font=F(12,bold=True), cursor="hand2",
            command=self._toggle_watch)
        self.watch_btn.pack(side="left", padx=(10,4))

        tk.Button(tb, text="📂  ファイル追加",
            bg="#1e40af", fg="white", relief="flat",
            padx=12, pady=5, font=F(11), cursor="hand2",
            command=self._add_files).pack(side="left", padx=4)

        tk.Button(tb, text="⚡  今すぐ全処理",
            bg="#78350f", fg="white", relief="flat",
            padx=12, pady=5, font=F(11), cursor="hand2",
            command=self._process_all).pack(side="left", padx=4)

        tk.Button(tb, text="🗑  リストクリア",
            bg="#4b5563", fg="white", relief="flat",
            padx=12, pady=5, font=F(11), cursor="hand2",
            command=self._clear_list).pack(side="left", padx=4)

        # セパレーター
        tk.Frame(tb, bg="#555555", width=1).pack(side="left", fill="y", padx=10, pady=3)

        tk.Label(tb, text="翻訳エンジン：", bg="#2d2d2d", fg="#cccccc", font=F(11)).pack(side="left")
        self.engine_var = tk.StringVar(value=self.cfg.get("engine","google"))
        eng_cb = ttk.Combobox(tb, textvariable=self.engine_var,
                              values=["google", "llm"],
                              state="readonly", width=10, font=F(11))
        eng_cb.pack(side="left", padx=(2,10))
        eng_cb.bind("<<ComboboxSelected>>", lambda e: self._cfg_save())

        tk.Label(tb, text="⚙", bg="#2d2d2d", fg="#aaaaaa", font=F(14),
                 cursor="hand2").pack(side="right", padx=10)

    # ── メインエリア ─────────────────────────────────
    def _body(self):
        pw = ttk.PanedWindow(self.root, orient="horizontal")
        pw.pack(fill="both", expand=True, padx=8, pady=8)

        # ── 左: ファイルキュー ──
        left = tk.Frame(pw, bg=APP_BG)
        pw.add(left, weight=2)

        hdr_l = tk.Frame(left, bg=APP_BG)
        hdr_l.pack(fill="x", pady=(0,4))
        tk.Label(hdr_l, text="📋  ファイルキュー", bg=APP_BG,
                 fg=APP_DARK, font=F(12,bold=True)).pack(side="left")
        tk.Label(hdr_l,
                 text=f"対応: TXT{'  Excel' if EXCEL_OK else ''}{'  Word' if WORD_OK else ''}",
                 bg=APP_BG, fg="#6b7280", font=F(10)).pack(side="right")

        cols = ("name","type","status","time")
        self.tree = ttk.Treeview(left, columns=cols, show="headings",
                                  selectmode="browse")
        self.tree.heading("name",   text="ファイル名")
        self.tree.heading("type",   text="種類")
        self.tree.heading("status", text="状態")
        self.tree.heading("time",   text="処理時刻")
        self.tree.column("name",   width=180, anchor="w")
        self.tree.column("type",   width=70,  anchor="center")
        self.tree.column("status", width=90,  anchor="center")
        self.tree.column("time",   width=80,  anchor="center")

        self.tree.tag_configure("waiting",    foreground="#6b7280")
        self.tree.tag_configure("processing", foreground="#1e40af", background="#eff6ff")
        self.tree.tag_configure("done",       foreground="#166534", background="#f0fdf4")
        self.tree.tag_configure("error",      foreground="#991b1b", background="#fff1f2")

        sb = ttk.Scrollbar(left, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # ── 右: 統計 + ログ ──
        right = tk.Frame(pw, bg=APP_BG)
        pw.add(right, weight=3)

        # 統計カード
        card_row = tk.Frame(right, bg=APP_BG)
        card_row.pack(fill="x", pady=(0,8))
        self.stat_labels = {}
        for key, label, color in [
            ("total", "合計",   "#1e40af"),
            ("done",  "完了",   "#166534"),
            ("error", "エラー", "#991b1b"),
        ]:
            card = tk.Frame(card_row, bg=APP_PANEL, bd=1, relief="solid")
            card.pack(side="left", expand=True, fill="x", padx=(0,6))
            tk.Label(card, text=label, bg=APP_PANEL, fg="#6b7280",
                     font=F(10), pady=4).pack()
            lbl = tk.Label(card, text="0", bg=APP_PANEL, fg=color,
                           font=F(26,bold=True))
            lbl.pack(pady=(0,6))
            self.stat_labels[key] = lbl

        # プログレスバー
        prog_f = tk.Frame(right, bg=APP_BG)
        prog_f.pack(fill="x", pady=(0,6))
        tk.Label(prog_f, text="進捗", bg=APP_BG, fg="#6b7280", font=F(10)).pack(anchor="w")
        self.progress = ttk.Progressbar(prog_f, mode="determinate", length=400)
        self.progress.pack(fill="x")
        self.prog_label = tk.Label(prog_f, text="", bg=APP_BG, fg="#6b7280", font=F(10))
        self.prog_label.pack(anchor="e")

        # ログ
        tk.Label(right, text="処理ログ", bg=APP_BG, fg=APP_DARK,
                 font=F(11,bold=True), anchor="w").pack(fill="x")
        self.log_box = scrolledtext.ScrolledText(
            right, font=(FONT_MONO,10),
            bg="#1e1b1b", fg="#d4d4d4",
            insertbackground="white", relief="flat", bd=0)
        self.log_box.pack(fill="both", expand=True, pady=(3,0))
        # ログカラータグ
        self.log_box.tag_config("info",    foreground="#d4d4d4")
        self.log_box.tag_config("ok",      foreground="#86efac")
        self.log_box.tag_config("warn",    foreground="#fde047")
        self.log_box.tag_config("err",     foreground="#fca5a5")
        self.log_box.tag_config("section", foreground="#67e8f9")

    # ── ステータスバー ───────────────────────────────
    def _statusbar(self):
        sb = tk.Frame(self.root, bg="#2d2d2d", pady=4)
        sb.pack(fill="x", side="bottom")
        self.status_var = tk.StringVar(value="⏸  停止中")
        tk.Label(sb, textvariable=self.status_var,
                 bg="#2d2d2d", fg="#cccccc", font=F(10)).pack(side="left", padx=10)
        self.folder_var = tk.StringVar(value=f"入力: {self.in_dir}  |  出力: {self.out_dir}")
        tk.Label(sb, textvariable=self.folder_var,
                 bg="#2d2d2d", fg="#888888", font=F(10)).pack(side="right", padx=10)

    # ── 設定ロード / セーブ ─────────────────────────
    def _cfg_load(self):
        d = {"engine":"google","llm_url":"http://localhost:11434",
             "llm_model":"llama3.2","llm_api":"ollama",
             "in_dir": str(BASE_DIR/"input"),
             "out_dir": str(BASE_DIR/"output")}
        if CONFIG_FILE.exists():
            try: d.update(json.loads(CONFIG_FILE.read_text(encoding="utf-8")))
            except: pass
        return d

    def _cfg_save(self):
        self.cfg["engine"] = self.engine_var.get()
        CONFIG_FILE.write_text(
            json.dumps(self.cfg, ensure_ascii=False, indent=2), encoding="utf-8")

    # ── 起動ログ ─────────────────────────────────────
    def _boot_log(self):
        self._log(f"翻訳ウォッチャー Pro 起動  [{platform.system()} / Python {sys.version.split()[0]}]", "section")
        missing = []
        if not GOOGLE_OK: missing.append("deep-translator")
        if not EXCEL_OK:  missing.append("openpyxl")
        if not WORD_OK:   missing.append("python-docx")
        if not WATCH_OK:  missing.append("watchdog")
        if missing:
            self._log(f"未インストール: {', '.join(missing)}", "warn")
            self._log("  pip install " + " ".join(missing), "warn")
        self._log(f"対応形式: {', '.join(sorted(EXTS))}", "info")
        self._log(f"入力フォルダ: {self.in_dir}", "info")
        self._log(f"出力フォルダ: {self.out_dir}", "info")
        self._log("準備完了", "ok")

    def _log(self, msg: str, tag: str = "info"):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_box.insert("end", f"[{ts}]  {msg}\n", tag)
        self.log_box.see("end")

    def _update_stats(self):
        for k, lbl in self.stat_labels.items():
            lbl.config(text=str(self.stats[k]))
        total = self.stats["total"]
        done  = self.stats["done"] + self.stats["error"]
        self.progress["maximum"] = max(total, 1)
        self.progress["value"]   = done
        self.prog_label.config(text=f"{done} / {total}")

    # ── 監視制御 ─────────────────────────────────────
    def _toggle_watch(self):
        if self.watching: self._stop_watch()
        else:             self._start_watch()

    def _start_watch(self):
        if not WATCH_OK:
            messagebox.showerror("エラー","watchdog が未インストールです。\npip install watchdog")
            return
        self._cfg_save()
        h = _Handler(self)
        self.observer = Observer()
        self.observer.schedule(h, str(self.in_dir), recursive=False)
        self.observer.start()
        self.watching = True
        self.watch_btn.config(text="■  監視 停止", bg="#991b1b")
        eng = "Google翻訳" if self.engine_var.get()=="google" else f"LLM({self.cfg.get('llm_model','')})"
        self.status_var.set(f"🟢  監視中  [{eng}]")
        self._log(f"監視開始 — {self.in_dir}", "ok")

    def _stop_watch(self):
        if self.observer:
            self.observer.stop(); self.observer.join(); self.observer = None
        self.watching = False
        self.watch_btn.config(text="▶  監視 開始", bg="#166534")
        self.status_var.set("⏸  停止中")
        self._log("監視停止", "warn")

    # ── ファイル追加 ─────────────────────────────────
    def _add_files(self):
        ft = [("対応ファイル", " ".join(f"*{e}" for e in sorted(EXTS))),
              ("テキスト","*.txt"),("All","*.*")]
        if EXCEL_OK: ft.insert(2,("Excel","*.xlsx *.xls"))
        if WORD_OK:  ft.insert(3,("Word","*.docx *.doc"))
        paths = filedialog.askopenfilenames(title="翻訳するファイルを選択", filetypes=ft)
        for p in paths:
            threading.Thread(target=self._process_file, args=(Path(p),), daemon=True).start()

    def _process_all(self):
        self._cfg_save()
        files = [f for f in self.in_dir.iterdir() if f.suffix.lower() in EXTS]
        if not files:
            self._log(f"input フォルダに対応ファイルがありません", "warn"); return
        for f in files:
            threading.Thread(target=self._process_file, args=(f,), daemon=True).start()

    def _clear_list(self):
        for item in self.tree.get_children(): self.tree.delete(item)
        self.stats = {"total":0,"done":0,"error":0}
        self._update_stats()

    # ── ファイル処理（振り分け）─────────────────────
    def _process_file(self, path: Path):
        ext = path.suffix.lower()
        iid = self._tree_add(path, ext)
        self.stats["total"] += 1
        self.root.after(0, self._update_stats)

        try:
            if ext == ".txt":                      self._do_txt(path, iid)
            elif ext in (".xlsx",".xls"):          self._do_excel(path, iid)
            elif ext in (".docx",".doc"):          self._do_word(path, iid)
            else:
                raise ValueError(f"非対応の形式: {ext}")
            self.stats["done"] += 1
            self.root.after(0, lambda: self._tree_set(iid,"done","✅ 完了"))
        except Exception as e:
            self.stats["error"] += 1
            self.root.after(0, lambda: self._tree_set(iid,"error",f"❌ {e}"))
            self.root.after(0, self._log, f"❌ {path.name}: {e}", "err")
        finally:
            self.root.after(0, self._update_stats)

    # ── TXT翻訳 ─────────────────────────────────────
    def _do_txt(self, path: Path, iid: str):
        self.root.after(0, self._log, f"━━ TXT: {path.name}", "section")
        text = _read(path)
        lines = text.splitlines()
        self.root.after(0, self._log, f"   {len(lines)} 行", "info")
        for lc, ln in LANGS.items():
            out = []
            for line in lines:
                out.append(self._tr(line, lc, ln) if line.strip() else "")
            dest = self.out_dir / f"{path.stem}_{lc}{path.suffix}"
            dest.write_text("\n".join(out), encoding="utf-8-sig")
            self.root.after(0, self._log, f"   ✓ {ln} → {dest.name}", "ok")

    # ── Excel翻訳 ────────────────────────────────────
    def _do_excel(self, path: Path, iid: str):
        if not EXCEL_OK: raise RuntimeError("openpyxl 未インストール")
        self.root.after(0, self._log, f"━━ Excel: {path.name}", "section")
        for lc, ln in LANGS.items():
            wb  = openpyxl.load_workbook(path)
            cnt = 0
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and cell.value.strip():
                            cell.value = self._tr(cell.value, lc, ln)
                            cnt += 1
            dest = self.out_dir / f"{path.stem}_{lc}{path.suffix}"
            wb.save(dest)
            self.root.after(0, self._log, f"   ✓ {ln} {cnt}セル → {dest.name}", "ok")

    # ── Word翻訳 ─────────────────────────────────────
    def _do_word(self, path: Path, iid: str):
        if not WORD_OK: raise RuntimeError("python-docx 未インストール")
        self.root.after(0, self._log, f"━━ Word: {path.name}", "section")
        for lc, ln in LANGS.items():
            doc = _Docx(path)
            cnt = 0
            # 本文段落
            for para in doc.paragraphs:
                if para.text.strip():
                    _replace_para(para, self._tr(para.text, lc, ln))
                    cnt += 1
            # テーブル
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                _replace_para(para, self._tr(para.text, lc, ln))
                                cnt += 1
            dest = self.out_dir / f"{path.stem}_{lc}{path.suffix}"
            doc.save(dest)
            self.root.after(0, self._log, f"   ✓ {ln} {cnt}段落 → {dest.name}", "ok")

    # ── 翻訳エンジン呼び出し ─────────────────────────
    def _tr(self, text: str, lc: str, ln: str) -> str:
        if not text.strip(): return text
        if self.engine_var.get() == "google":
            return self._google(text, lc)
        return self._llm(text, ln)

    def _google(self, text: str, lc: str) -> str:
        if not GOOGLE_OK: return text
        try: return GoogleTranslator(source="ja", target=lc).translate(text)
        except: return text

    def _llm(self, text: str, ln: str) -> str:
        url   = self.cfg.get("llm_url","http://localhost:11434").rstrip("/")
        model = self.cfg.get("llm_model","llama3.2")
        api   = self.cfg.get("llm_api","ollama")
        prompt = f"以下の日本語テキストを{ln}に翻訳してください。翻訳文のみ返してください。\n\n{text}"
        try:
            if "openai" in api:
                body = json.dumps({"model":model,"messages":[{"role":"user","content":prompt}],"temperature":0.1}).encode()
                req  = urllib.request.Request(f"{url}/v1/chat/completions", data=body,
                                              headers={"Content-Type":"application/json"}, method="POST")
                with urllib.request.urlopen(req, timeout=60) as r:
                    return json.loads(r.read())["choices"][0]["message"]["content"].strip()
            else:
                body = json.dumps({"model":model,"prompt":prompt,"stream":False}).encode()
                req  = urllib.request.Request(f"{url}/api/generate", data=body,
                                              headers={"Content-Type":"application/json"}, method="POST")
                with urllib.request.urlopen(req, timeout=60) as r:
                    return json.loads(r.read()).get("response", text).strip()
        except: return text

    # ── Treeview操作 ─────────────────────────────────
    def _tree_add(self, path: Path, ext: str) -> str:
        type_label = {"txt":"TXT","xlsx":"Excel","xls":"Excel",
                      "docx":"Word","doc":"Word"}.get(ext.lstrip("."),"FILE")
        ts  = datetime.now().strftime("%H:%M")
        iid = f"{id(path)}_{time.time()}"
        self.root.after(0, lambda: self.tree.insert(
            "", "end", iid,
            values=(path.name, type_label, "⏳ 処理中", ts),
            tags=("processing",)))
        return iid

    def _tree_set(self, iid: str, tag: str, status: str):
        try:
            vals = (self.tree.set(iid,"name"),
                    self.tree.set(iid,"type"),
                    status,
                    self.tree.set(iid,"time"))
            self.tree.item(iid, values=vals, tags=(tag,))
        except: pass

    # ── 設定ダイアログ ───────────────────────────────
    def _open_settings(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("設定")
        dlg.geometry("480x400")
        dlg.resizable(False, False)
        dlg.configure(bg=APP_BG)
        dlg.grab_set()

        nb = ttk.Notebook(dlg)
        nb.pack(fill="both", expand=True, padx=12, pady=12)

        # ── フォルダタブ ──
        tab1 = tk.Frame(nb, bg=APP_BG, padx=16, pady=16)
        nb.add(tab1, text="フォルダ")

        in_var  = tk.StringVar(value=str(self.in_dir))
        out_var = tk.StringVar(value=str(self.out_dir))

        for label, var, setter in [
            ("入力フォルダ", in_var,  lambda: in_var.set(filedialog.askdirectory() or in_var.get())),
            ("出力フォルダ", out_var, lambda: out_var.set(filedialog.askdirectory() or out_var.get())),
        ]:
            tk.Label(tab1, text=label, bg=APP_BG, font=F(11,bold=True), anchor="w").pack(fill="x")
            r = tk.Frame(tab1, bg=APP_BG); r.pack(fill="x", pady=(2,10))
            tk.Entry(r, textvariable=var, font=F(10), width=42).pack(side="left")
            tk.Button(r, text="参照", bg=APP_RED, fg="white", relief="flat", padx=8,
                      command=setter).pack(side="left", padx=4)

        # ── LLMタブ ──
        tab2 = tk.Frame(nb, bg=APP_BG, padx=16, pady=16)
        nb.add(tab2, text="LLM設定")

        fields = [
            ("API種別",  "llm_api",   ["ollama","openai互換"]),
            ("URL",      "llm_url",   None),
            ("モデル名", "llm_model", None),
        ]
        vars2 = {}
        for label, key, choices in fields:
            tk.Label(tab2, text=label, bg=APP_BG, font=F(11), anchor="w").pack(fill="x")
            v = tk.StringVar(value=self.cfg.get(key,""))
            vars2[key] = v
            if choices:
                ttk.Combobox(tab2, textvariable=v, values=choices, state="readonly", width=30).pack(anchor="w", pady=(0,8))
            else:
                tk.Entry(tab2, textvariable=v, font=F(11), width=38).pack(anchor="w", pady=(0,8))

        # 保存ボタン
        def save():
            self.in_dir  = Path(in_var.get());  self.in_dir.mkdir(exist_ok=True)
            self.out_dir = Path(out_var.get()); self.out_dir.mkdir(exist_ok=True)
            self.cfg["in_dir"]  = str(self.in_dir)
            self.cfg["out_dir"] = str(self.out_dir)
            for k, v in vars2.items(): self.cfg[k] = v.get()
            self._cfg_save()
            self.folder_var.set(f"入力: {self.in_dir}  |  出力: {self.out_dir}")
            self._log("設定を保存しました", "ok")
            dlg.destroy()

        tk.Button(dlg, text="保存して閉じる", bg=APP_RED, fg="white",
                  relief="flat", padx=16, pady=8, font=F(12,bold=True),
                  command=save).pack(pady=8)

    # ── ヘルプ ──────────────────────────────────────
    def _show_help(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("使い方")
        dlg.geometry("500x380")
        dlg.configure(bg=APP_BG)
        dlg.grab_set()
        txt = scrolledtext.ScrolledText(dlg, font=F(11), bg=APP_PANEL, relief="flat")
        txt.pack(fill="both", expand=True, padx=12, pady=12)
        txt.insert("end", """【翻訳ウォッチャー Pro  使い方】

■ 基本の流れ
  1. ツールバーの「翻訳エンジン」で Google翻訳 / LLM を選ぶ
  2. 「監視 開始」を押す
  3. input フォルダにファイルを入れると自動翻訳される
  4. output フォルダに翻訳済みファイルが保存される

■ 対応ファイル形式
  ・テキスト (.txt)
  ・Excel (.xlsx / .xls)   ※ openpyxl が必要
  ・Word  (.docx / .doc)   ※ python-docx が必要

■ 出力ファイル名
  例）作業指示.xlsx
      → 作業指示_vi.xlsx    (ベトナム語)
      → 作業指示_zh-CN.xlsx (中国語)

■ ライブラリのインストール
  pip install deep-translator openpyxl python-docx watchdog

■ 翻訳言語の変更
  翻訳ウォッチャー.py の LANGS 辞書を編集してください。
""")
        txt.config(state="disabled")

    def on_close(self):
        self._stop_watch()
        self._cfg_save()
        self.root.destroy()


# ── フォルダ監視 ──────────────────────────────────────
class _Handler(FileSystemEventHandler):
    def __init__(self, app):
        self.app  = app
        self._seen: set = set()

    def on_created(self, event):
        path = Path(event.src_path)
        if not event.is_directory and path.suffix.lower() in EXTS:
            if str(path) not in self._seen:
                self._seen.add(str(path))
                time.sleep(0.8)
                threading.Thread(target=self.app._process_file, args=(path,), daemon=True).start()


# ── ユーティリティ ────────────────────────────────────
def _read(path: Path) -> str:
    for enc in ("utf-8-sig","utf-8","shift-jis","cp932"):
        try: return path.read_text(encoding=enc)
        except: pass
    return ""

def _open(path: Path):
    if IS_WIN: os.startfile(str(path))
    else:
        import subprocess; subprocess.run(["open", str(path)])

def _replace_para(para, new_text: str):
    """Word段落のテキストを置換（書式を1つ目のrunに集約）"""
    if not para.runs: return
    para.runs[0].text = new_text
    for run in para.runs[1:]: run.text = ""


# ─────────────────────────────────────────────────────
if __name__ == "__main__":
    root = tk.Tk()
    app  = App(root)
    root.protocol("WM_DELETE_WINDOW", app.on_close)
    root.mainloop()
